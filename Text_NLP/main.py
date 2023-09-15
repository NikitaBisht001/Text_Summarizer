import os
from os.path import join, splitext
import docx
import para as para
import spacy
from PyPDF2 import PdfReader
from spacy.lang.en.stop_words import STOP_WORDS
from string import punctuation
import imaplib
import pytesseract
from pytesseract import image_to_string
from PIL import Image
from io import BytesIO
import fitz
from pdf2image import convert_from_path


current_dir = os.getcwd()
#print("Current Directory:", current_dir)
# Change the current working directory to the desired folder
os.chdir(r'D:\test')
# Access a specific folder within the current directory
target_folder = join(current_dir, r'D:\test')
print("Target Folder Path:", target_folder)
# List files and subdirectories within a folder
file_list = os.listdir(target_folder)
print("Files in Target Folder:", file_list)
stopwords = list(STOP_WORDS)
nlp = spacy.load('en_core_web_sm')
content = ""
pytesseract.pytesseract.tesseract_cmd = r'D:\tesseract\tesseract.exe'
for i in file_list:
    s = r"D:\test\{}".format(i)
    def get_file_extension(s):
        return splitext(s)[1][1:]  # Remove the leading dot
    # Example usage:
    extension = get_file_extension(s)
    # print("File extension:", extension)
    # Load the Word document

    if (extension == 'docx'):
        doct = docx.Document(s)
        content = "\n".join([para.text for para in doct.paragraphs])
        # Read the content of the document
        # content = "/n" .join
        # for para in doc.paragraphs:
        # small = doc.text
        # small = doc.text

    elif(extension == 'pdf'):
        def check_pdf_content(pdf_path):
            has_text = False
            has_images = False
            with open(pdf_path, 'rb') as pdf_file:
                pdf_reader = PdfReader(pdf_file)
                # Check if the PDF contains text
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():  # Check if the extracted text is not empty
                        has_text = True
                        break
                # Check if the PDF contains images
                if '/XObject' in page['/Resources']:
                    xObject = page['/Resources']['/XObject'].get_object()
                    if xObject:
                        for obj in xObject:
                            if xObject[obj]['/Subtype'] == '/Image':
                                has_images = True
                                break
            return has_text, has_images


        if __name__ == "__main__":
            text_present, images_present = check_pdf_content(s)

            if text_present:
                # If the page has text, extract it directly
                with open(s, 'rb') as pdf_file:
                    # Create a PDF reader object
                    pdf = PdfReader(pdf_file)
                    # Get the number of pages in the PDF
                    num_pages = len(pdf.pages)
                    # Extract text from all pages and store it in the 'text' variable
                    content = ''
                    for page_num in range(num_pages):
                        page = pdf.pages[page_num]
                        content += page.extract_text()
                # print("The PDF contains text.")

            elif images_present:
                # Open the PDF file and read its content using PyPDF2
                with open(s, 'rb') as pdf_file:
                    pdf_reader = PdfReader(pdf_file)
                    # Iterate through all pages
                    for page_num in range(len(pdf_reader.pages)):
                        # Convert the page to an image using pdf2image
                        images = convert_from_path(s, first_page=page_num + 1, last_page=page_num + 1, poppler_path=r'D:\poppler-23.07.0\Library\bin')
                        # Save each image to the output folder
                        for idx, image in enumerate(images):
                            image_path = join(f"page{page_num + 1}_image{idx + 1}.jpg")
                            img = image.save(image_path, 'JPEG')
                            extracted_text = pytesseract.image_to_string(Image.open(image_path))
                            content += extracted_text + "\n"

    doct = nlp(content)
    tokens = [token.text for token in doct]
    #print(tokens)
    punctuation = punctuation + '\n'
    punctuation
    word_frequencies = {}
    for word in doct:
        if word.text.lower() not in stopwords:
            if word.text.lower() not in punctuation:
                if word.text not in word_frequencies.keys():
                    word_frequencies[word.text] = 1
                else:
                    word_frequencies[word.text] += 1
                #print(word_frequencies)
                max_frequency = max(word_frequencies.values())
                max_frequency
    for word in word_frequencies.keys():
        word_frequencies[word] = word_frequencies[word] / max_frequency
        #print(word_frequencies)
    sentence_tokens = [sent for sent in doct.sents]
    print(sentence_tokens)
    sentence_scores = {}
    for sent in sentence_tokens:
        for word in sent:
            if word.text.lower() in word_frequencies.keys():
                if sent not in sentence_scores.keys():
                    sentence_scores[sent] = word_frequencies[word.text.lower()]
                else:
                    sentence_scores[sent] += word_frequencies[word.text.lower()]
                sentence_scores

    from heapq import nlargest

    select_length = int(len(sentence_tokens) * 0.3)
    select_length
    summary = nlargest(select_length, sentence_scores, key=sentence_scores.get)
    summary
    final_summary = [word.text for word in summary]
    summary = ' '.join(final_summary)
    print(summary)
    print(len(summary))
    print(len(content))

    

